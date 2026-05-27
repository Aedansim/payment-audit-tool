from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AMOUNT_COL = 'Payment Voucher Amount (SGD, Excluding GST)'

# Colour palette (RGB tuples)
NAVY  = RGBColor(0x1F, 0x38, 0x64)
BLUE  = RGBColor(0x2E, 0x75, 0xB6)
GREY  = RGBColor(0x60, 0x60, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED   = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x70, 0xAD, 0x47)


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.name = 'Times New Roman'
    return p


def _body(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size + 2)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = GREY
    return p


def _bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size + 2)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = GREY
    return p


def _coloured_para(doc, label, value, colour=NAVY, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(size + 2)
    r1.font.name = 'Times New Roman'
    r1.font.color.rgb = NAVY
    r2 = p.add_run(str(value))
    r2.font.size = Pt(size + 2)
    r2.font.name = 'Times New Roman'
    r2.font.color.rgb = colour
    return p


def _shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    w, h = section.page_height, section.page_width
    section.page_width = w
    section.page_height = h
    section.left_margin  = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin   = Cm(1.5)
    section.bottom_margin = Cm(1.5)


# ---------------------------------------------------------------------------
# Scope and Limitations — rendered first, before Executive Summary
# ---------------------------------------------------------------------------

def _page_caveats(doc):
    _heading(doc, "Scope and Limitations", level=1)
    _body(doc,
          "The following limitations should be understood before acting on the tool's output:",
          size=10)
    caveats = [
        "Output alone does not constitute proof of fraud or error — a high voucher score indicates a statistically unusual "
        "transaction that warrants attention; it is not evidence of fraud or error. Professional "
        "judgement is required for all selected vouchers. Transactions not flagged should not be "
        "assumed free from irregularities, as sophisticated anomalies that closely mimic normal "
        "payment patterns may go undetected.",
        "Line-item scope — the tool scores individual transaction lines, not total voucher amounts. "
        "A large voucher split across many small lines of normal individual amounts may not score "
        "highly even if the total is anomalous. Auditors should review total voucher values "
        "alongside individual line scores.",
        "Pre-calibrated weights — component weights and rule thresholds are calibrated for typical "
        "corporate payment datasets. Unusual compositions (e.g. predominantly recurring payments, "
        "narrow amount bands) may require recalibration. Weights can be overridden via "
        "'sample_selector.WEIGHTS' before calling select_samples().",
        "Declared component weights are approximate — the five weights describe intended relative "
        "importance, not precisely isolated statistical contributions. Features shared across "
        "components (amount z-scores and rule flags feed both their dedicated scoring components "
        "and Isolation Forest/LOF) carry marginally more effective influence than their labelled "
        "weight suggests. This does not affect the relative voucher rankings in practice.",
    ]
    for caveat in caveats:
        _bullet(doc, caveat, size=10)


# ---------------------------------------------------------------------------
# Page 1 — Executive Summary
# ---------------------------------------------------------------------------

def _page1(doc, df, df_vouchers, selected_vouchers, benford_stats):
    _heading(doc, "Executive Summary", level=1)

    _body(doc,
          "The payment vouchers in this report were selected using a composite risk score that "
          "combines the outputs of five independent analytical methods — machine learning anomaly "
          "detection (two methods), statistical z-score analysis, rule-based forensic checks, and Benford's Law "
          "deviation analysis — into a single ranked score per voucher. Vouchers are stratified "
          "into HIGH, MEDIUM, and LOW risk tiers based on their score percentile, and the audit "
          "sample prioritises higher-risk tiers. Diversity controls are applied to prevent any "
          "single vendor or payment pattern from dominating the sample; the specific criteria are "
          "described in the Methodology section of this report. The selected samples are intended "
          "as risk-based suggestions to guide audit focus. Auditors should exercise professional "
          "judgement in determining which payments to proceed with for further testing.",
          size=10)
    doc.add_paragraph()

    _body(doc,
          "The dataset overview, the summary of findings, and all supporting charts are presented "
          "in the accompanying Excel workbook. Refer to the 'Summary' tab for the dataset overview "
          "and summary of findings, and to the 'Analytical Charts' tab for the supporting charts.",
          size=10)
    doc.add_paragraph()

    _body(doc,
          f"Each of the {len(selected_vouchers)} selected payment vouchers is accompanied by "
          "documented reason codes that explain which analytical signals drove its inclusion, "
          "helping auditors prioritise their review approach and focus testing on the most "
          "material risk indicators. Scoring is performed at individual line-item level and "
          "rolled up to payment voucher level, so each entry in the sample corresponds to a "
          "complete voucher that can be physically retrieved for examination. Full details — "
          "including line-level scores, flags, and reason codes — are available in the "
          "'Selected Vouchers' and 'Voucher Line Detail' tabs in the accompanying Excel workbook.",
          size=10)


# ---------------------------------------------------------------------------
# Page 2 — Methodology
# ---------------------------------------------------------------------------

def _page2(doc, excluded_count=0):
    _heading(doc, "Methodology — How the Tool Works", level=1)

    _body(doc,
          "This section provides a full technical account of the analytical process applied to "
          "the payment data — from initial feature extraction through to final audit sample "
          "selection — at a level of detail sufficient for independent verification, methodology "
          "recalibration, or peer review. Readers who require only a high-level understanding of "
          "the selection approach are referred to the Executive Summary. Each stage is described "
          "in full below.",
          size=10)
    doc.add_paragraph()

    # ---- Stage 1 ----
    _heading(doc, "Stage 1 — Feature Engineering", level=2)
    _body(doc,
          "Feature engineering transforms raw payment transaction data into structured signals "
          "that the machine learning models and rule-based checks can evaluate. Without this step, "
          "the models would have no meaningful way to distinguish between a normal payment and an "
          "anomalous one — raw fields such as vendor name, invoice date, and payment amount carry "
          "limited signals on their own. By computing derived features such as how much a payment "
          "deviates from a vendor's historical average, whether the amount falls suspiciously close "
          "to an approval threshold, or whether the invoice numbering follows a sequential pattern "
          "on the same date, the tool surfaces patterns that would otherwise require manual "
          "inspection of thousands of individual transactions.",
          size=10)
    _body(doc,
          "A complete list of features computed in this stage, including their definitions, data "
          "types, and usage across the scoring components, is provided in the Features Reference "
          "Table in the later section of this report.",
          size=10)
    doc.add_paragraph()

    # ---- Stage 2 ----
    _heading(doc, "Stage 2 — Five Independent Analytical Methods", level=2)
    _body(doc,
          "Each transaction line is independently assessed by five methods. Using multiple independent "
          "methods reduces both false positives (legitimate transactions wrongly flagged) and false "
          "negatives (genuine anomalies missed). No single method is relied upon alone.",
          size=10)
    doc.add_paragraph()

    _heading(doc, "1. Benford's Law", level=2)
    _body(doc,
          "In any large collection of naturally occurring financial amounts, approximately 30% start "
          "with digit 1, 17% with 2, 12% with 3, declining to 5% for digit 9. Significant deviation "
          "from this pattern may indicate amounts were manually entered, rounded, or constructed. "
          "Deviation is measured using the Mean Absolute Deviation (MAD) — with Non-Conformity "
          "defined as MAD > 0.015 (Nigrini, 2012) — and a chi-square significance test.",
          size=10)
    _body(doc,
          "Caveat: Benford's Law is most reliable for large datasets (ideally > 1,000 non-recurring "
          "transactions). Smaller datasets or narrow amount ranges produce less stable results. "
          "The chi-square test is very sensitive for large datasets and may flag minor deviations "
          "as statistically significant even when they are not practically meaningful — MAD is the "
          "primary practical measure. Recurring payments (monthly, quarterly, semi-annual, annual cycles) are detected "
          "and tagged separately. They are excluded from Benford's Law analysis because their fixed "
          "amounts naturally deviate from Benford's expected distribution without being suspicious. ",
          italic=True, size=10)
    doc.add_paragraph()

    _heading(doc, "2. Isolation Forest (Machine Learning)", level=2)
    _body(doc,
          "An unsupervised machine learning model that detects anomalies by repeatedly splitting "
          "the data using random rules until each transaction is isolated. Transactions genuinely "
          "different from the rest require fewer splits to isolate — they are unusual in many "
          "dimensions simultaneously. The model evaluates all engineered features together: amount, "
          "processing time, date attributes, payee type, and vendor patterns.",
          size=10)
    _body(doc,
          "Caveat: Being unsupervised, the model identifies outliers relative to the current dataset. "
          "If the dataset contains pervasive irregularities, they may appear normal relative to each "
          "other and not be flagged. The model is most effective when the majority of transactions "
          "are legitimate.",
          italic=True, size=10)
    doc.add_paragraph()

    _heading(doc, "3. Local Outlier Factor — LOF (Machine Learning)", level=2)
    _body(doc,
          "LOF compares each transaction to its nearest neighbours — the most similar transactions "
          "by amount, vendor, and timing. A payment may look ordinary across the full dataset but "
          "be highly anomalous among its vendor peers. For example, a $50,000 payment to a vendor "
          "whose typical invoices are around $2,000 would score very highly even if $50,000 appears "
          "elsewhere in the dataset. This context-sensitivity makes LOF particularly effective for "
          "catching inflated invoices or payments to unusual recipients.",
          size=10)
    _body(doc,
          "Caveat: Same unsupervised limitation as Isolation Forest applies.",
          italic=True, size=10)
    doc.add_paragraph()

    _heading(doc, "4. Statistical Z-Score Analysis", level=2)
    _body(doc,
          "For each vendor and each cost centre, the average payment amount and standard deviation "
          "are computed across all transactions in the dataset. Payments more than 2 standard "
          "deviations above their group average are flagged — a threshold derived from the normal "
          "distribution, where ±2 standard deviations encompasses approximately 95% of values, "
          "leaving the upper 2.5% as statistical outliers. The 2-standard-deviation threshold is a widely "
          "applied convention in quantitative analysis. This approach is consistent with the objective of analytical "
          "procedures which requires auditors to identify and investigate significant fluctuations "
          "or relationships that are inconsistent with other relevant information or that differ "
          "from expected values.",
          size=10)
    _body(doc,
          "Caveat: Auditors should apply professional judgement in assessing whether flagged amounts "
          "are significant in context, noting that payment amounts may follow a skewed rather than "
          "normal distribution, which means the proportion flagged may differ from the theoretical 2.5%.",
          italic=True, size=10)
    doc.add_paragraph()

    _heading(doc, "5. Rule-Based Flags", level=2)
    _body(doc,
          "Nine binary rules derived from established forensic audit practice. Each triggers a "
          "flag (1) or not (0) per transaction line:",
          size=10)
    rules = [
        "Round number — amount divisible by 100. Round number amounts may warrant attention as "
        "fabricated or manually chosen amounts sometimes exhibit round number bias, where "
        "individuals select psychologically convenient figures rather than amounts arising from "
        "genuine invoices (Nigrini, 2012; ACFE Fraud Examiners Manual).",
        "Weekend payment — voucher accounting date falls on a Saturday or Sunday. "
        "Payments processed outside business hours may bypass the normal multi-person review process.",
        "Near approval threshold — within 5% below SGD 1K / 5K / 10K / 50K / 100K. Known as "
        "'structuring' in forensic accounting — deliberately keeping amounts below authorisation "
        "thresholds to avoid triggering higher-level approval.",
        "Individual payee — Vendor ID matches the Singapore NRIC/FIN format (one letter, 7 digits, "
        "one letter). Payments to individuals carry higher inherent risk versus registered businesses.",
        "Irregular repeated amount — same amount paid to the same vendor more than twice with no "
        "detected regular monthly/quarterly/annual schedule. May indicate split or duplicate payments.",
        "Duplicate payment — the same invoice number, invoice date, vendor, and amount appears "
        "across more than one distinct payment voucher. Indicates a potential double payment of "
        "the same invoice.",
        "Reversal or credit note — payment amount is negative. Reversals and credit notes are "
        "legitimate but warrant review, particularly when paired with other risk signals on the "
        "corresponding original payment.",
        "Split purchase risk — the same vendor has two or more invoices on the same invoice date "
        "with alphanumerically sequential invoice number suffixes (e.g. INV-1001, INV-1002) AND "
        "the group total falls within 5% below $6,000 or $90,000 (i.e. $5,700–<$6,000 or "
        "$85,500–<$90,000). May indicate a single purchase deliberately split across multiple "
        "invoices to avoid triggering a higher-level approval threshold.",
        "Transposed amount — the payment amount differs from another transaction to the same vendor "
        "with the same description by exactly one pair of transposed digits in the whole-dollar "
        "portion of the amount, with cents ignored (e.g. SGD 4,800 vs SGD 8,400). May indicate a "
        "keying error resulting in significant over- or under-payment if left undetected.",
    ]
    for rule in rules:
        _bullet(doc, rule, size=10)
    doc.add_paragraph()

    # ---- Stage 3 — Scoring formulas ----
    _heading(doc, "Stage 3 — Scoring Formulas and Weight Rationale", level=2)

    _body(doc, "Line-Level Composite Risk Score", bold=True, size=10)
    _body(doc,
          "Each of the five methods produces a score between 0 and 1, where 0 means most normal and 1 means most anomalous. "
          "These are combined into a single risk score using fixed weights:",
          size=10)
    p = _body(doc,
              "    risk_score  =  0.30 × IF  +  0.25 × LOF  +  0.25 × Z-score"
              "  +  0.15 × rule_flags  +  0.05 × Benford",
              bold=True, size=8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    _body(doc,
          "The rule flags score is the fraction of the 9 binary rules triggered for that line "
          "(e.g. 2 rules triggered = 2/9 = 0.22). The Benford score is normalised relative to "
          "the maximum Benford deviation in the dataset. The Z-score signal is the larger of the "
          "vendor z-score and cost centre z-score, min-max normalised to [0, 1] across all lines.",
          size=10)
    doc.add_paragraph()

    _body(doc, "Weight Rationale", bold=True, size=10)
    doc.add_paragraph()

    weight_rows = [
        ("Isolation Forest", "30%",
         "Primary ML signal; highest weight because it evaluates all features simultaneously "
         "and captures complex multi-dimensional patterns invisible to individual rules or "
         "statistics alone."),
        ("Local Outlier Factor", "25%",
         "Context-sensitive complement to Isolation Forest. Peer-group benchmarking reduces false "
         "positives by comparing each transaction to its most similar counterparts rather than "
         "the full dataset."),
        ("Z-Score Analysis", "25%",
         "Transparent and directly auditable. "
         "Higher weight because it is statistically rigorous and independently defensible."),
        ("Rule-Based Flags", "15%",
         "Directly encodes established forensic audit heuristics. Lower weight because rules are "
         "binary (on/off) and each has known limitations; their primary value is confirming and "
         "explaining signals raised by the other methods."),
        ("Benford's Law", "5%",
         "Supplementary signal only. Powerful at the dataset level but noisy at the individual "
         "transaction level. Low weight prevents Benford deviation alone from driving selection. "
         "Further suppressed when all other signals are below average (see rule below)."),
    ]

    tbl = doc.add_table(rows=1 + len(weight_rows), cols=3)
    tbl.style = 'Table Grid'
    col_widths_wt = [Inches(1.5), Inches(0.6), Inches(4.1)]
    hdr = tbl.rows[0]
    for i, label in enumerate(["Method", "Weight", "Basis for Weight Assignment"]):
        cell = hdr.cells[i]
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _shade_cell(cell, "1F3864")
        cell.width = col_widths_wt[i]

    for row_idx, (method, weight, rationale) in enumerate(weight_rows, start=1):
        row = tbl.rows[row_idx]
        shade = "F2F6FC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, (value, width) in enumerate(zip([method, weight, rationale], col_widths_wt)):
            cell = row.cells[col_idx]
            cell.text = value
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _shade_cell(cell, shade)
            cell.width = width

    doc.add_paragraph()

    _body(doc, "Benford Suppression Rule", bold=True, size=10)
    _body(doc,
          "If a transaction's Isolation Forest, LOF, Z-score, and rule flags scores are ALL below "
          "their respective dataset medians — meaning the transaction shows no elevated risk on any "
          "other signal — its Benford contribution is zeroed out entirely. This prevents Benford "
          "deviation alone from selecting a transaction. Benford evidence is only counted when at "
          "least one other signal is also elevated.",
          size=10)
    doc.add_paragraph()

    # ---- Stage 4 — Voucher rollup ----
    _heading(doc, "Stage 4 — Voucher-Level Rollup", level=2)
    _body(doc,
          "Individual scored lines are grouped by Voucher ID — the document auditors physically "
          "pull — rather than by invoice number. The voucher score formula is:",
          size=10)
    p = _body(doc,
              "    voucher_score (multi-line)  =  0.60 × max_line_score"
              "  +  0.25 × mean_line_score  +  0.15 × flag_density",
              bold=True, size=8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = _body(doc,
              "    voucher_score (single-line)  =  line risk_score  (no rollup needed)",
              bold=True, size=8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    _body(doc,
          "Flag density = total rule flags triggered across all lines in the voucher ÷ "
          "(9 flag types × number of lines). The 60/25/15 split reflects that audit significance "
          "is primarily driven by the worst line in the voucher, moderated by whether other lines "
          "are also elevated, and supplemented by the breadth of rule flag coverage. For multi-line "
          "vouchers, reason codes in the output are prefixed with [Account Code] so auditors can "
          "identify exactly which line triggered each flag.",
          size=10)
    doc.add_paragraph()

    # ---- ML Consensus ----
    _heading(doc, "ML Consensus Flag", level=2)
    _body(doc,
          "Each transaction line receives an ML Consensus count: the number of the three ML-based "
          "methods that independently classify that line as anomalous using each model's own "
          "boundary. Isolation Forest and LOF use sklearn's predict() method at "
          "contamination=0.05, which flags the top 5% of lines as anomalous per model. "
          "Z-score flags lines where the maximum absolute z-score exceeds 2.0 (2 standard "
          "deviations). A voucher is marked 'ML Consensus = Yes' if any of its lines is "
          "classified as anomalous by 2 or more of the three methods simultaneously.",
          size=10)
    _body(doc,
          "The ML Consensus flag does not alter the composite score — it is a corroborating "
          "indicator shown in the Excel output. When multiple independent methods agree that a "
          "transaction is anomalous, the probability of a true anomaly is materially higher than "
          "when only one method flags it.",
          size=10)
    doc.add_paragraph()

    # ---- Potential FY split purchases (review aid, not scored) ----
    _heading(doc, "Potential Fiscal Year Split Purchases", level=2)
    _body(doc,
          "The tool identifies potential split purchases within a fiscal year. A fiscal year runs "
          "from 1 April to 31 March of the following year. This check looks for multiple payments of a "
          "similar nature made to the same vendor within the same fiscal year where the combined total "
          "exceeds SGD 6,000 — the approval threshold for small value purchases. Such a pattern may "
          "indicate that a single procurement need has been deliberately divided into smaller payments, "
          "each kept below the threshold, to avoid the approval process that would apply to the "
          "combined amount.",
          size=10)
    _body(doc,
          "Payments are grouped by vendor, fiscal year, and a normalised payment description (which "
          "removes numbers and punctuation so that payments of the same general nature are grouped "
          "together regardless of invoice-specific detail). Where the positive-amount total of such a "
          "group exceeds SGD 6,000, every positive-amount transaction in the group is identified as a "
          "potential fiscal year split purchase. The fiscal year is determined using the Voucher "
          "Accounting Date. Where a payment's description consists of a structured reference code "
          "rather than descriptive text (for example a document or contract reference), the full "
          "reference is retained as the grouping key so the payment is not excluded from the check; "
          "in such cases payments are grouped only where they share the identical reference. A "
          "consequence is that vendors whose payments each carry a distinct reference code will not be "
          "grouped by this check and should be considered for separate manual review.",
          size=10)
    _body(doc,
          "This feature does NOT influence the composite risk score or the machine learning models. "
          "It is deliberately excluded from scoring because legitimate recurring payments — such as "
          "monthly subscriptions, retainers, or regular supplies of the same nature — would otherwise "
          "generate a high rate of false positives. Including such a signal in the scoring models would "
          "risk distorting the risk ranking. Instead, this feature operates purely as a standalone "
          "review aid: all identified cases are listed in the dedicated 'FY Split Purchase' tab of the "
          "accompanying Excel workbook, and a note is added to the reason codes of affected "
          "transactions. Auditors should review the listed groups directly and apply professional "
          "judgement to distinguish genuine procurement splitting from legitimate recurring payments, "
          "taking into account the vendor relationship, the existence of contracts or purchase orders, "
          "and the nature of the goods or services.",
          size=10)
    doc.add_paragraph()

    # ---- Risk tiers and selection ----
    _heading(doc, "Risk Tier Assignment and Sample Selection", level=2)
    _body(doc,
          "After all voucher scores are computed, tiers are assigned based on where each voucher ranks within the dataset: "
          "the top 5% of scores are flagged HIGH, the next 15% MEDIUM, and the remaining 80% LOW. "
          "Percentile-based tiers ensure the tool adapts to any dataset — HIGH always covers the "
          "most anomalous 5% regardless of absolute score values, which vary by dataset size and "
          "composition.",
          size=10)
    doc.add_paragraph()
    _body(doc,
          "The audit sample is drawn from all three tiers using a stratified approach. All "
          "HIGH-tier vouchers are included as mandatory selections. Remaining slots are filled "
          "proportionally from MEDIUM and LOW tiers, ensuring coverage of elevated-risk items "
          "while also providing a baseline of lower-risk vouchers against which audit findings "
          "can be contextualised.",
          size=10)
    doc.add_paragraph()
    _body(doc,
          "Within each vendor, near-duplicate vouchers are identified and consolidated before "
          "the sample is finalised. For any vendor with two or more selected vouchers, pairwise "
          "Jaccard token-overlap similarity is computed on the voucher line descriptions. If two "
          "vouchers share more than 70% of their description tokens, the lower-scoring voucher is "
          "replaced by the next-highest-scoring unselected voucher — drawn from any vendor — whose "
          "description does not introduce a new near-duplicate. This prevents the sample from being "
          "dominated by cosmetically similar payments when other analytically distinct, "
          "higher-scoring vouchers are available.",
          size=10)
    doc.add_paragraph()
    _body(doc,
          "A vendor cap of two vouchers per Vendor ID is then applied. If more than two vouchers "
          "from the same vendor remain selected after the similarity step, only the two "
          "highest-scoring are retained; excess vouchers are replaced by the next-best unselected "
          "alternatives that satisfy both the vendor cap and the similarity check. This ensures "
          "the sample spans a broad range of vendors and reduces the risk that audit attention "
          "is concentrated on a single supplier relationship.",
          size=10)
    doc.add_paragraph()
    _body(doc,
          "Vendors listed in the organisation's 'Excluded vendors' file (matched by UEN, which "
          "corresponds to Vendor ID) are de-prioritised in the sample selection process. This list "
          "is maintained by the organisation and typically includes government agencies and other "
          "vendors considered out of scope. Such vendors are excluded from the HIGH and MEDIUM "
          "selection tiers and will only appear in the audit sample if insufficient non-excluded "
          f"vouchers exist in the lower risk tier. {excluded_count} vendor UEN(s) were loaded from "
          "the Excluded vendors file for this run.",
          size=10)
    doc.add_paragraph()



# ---------------------------------------------------------------------------
# Feature Reference Table (landscape)
# ---------------------------------------------------------------------------

# Columns: Feature | What It Measures | Threshold for Flagging | ML Models | Why It Matters
# ML Models shows which of the three ML scoring components each feature feeds into:
#   IF  = Isolation Forest   LOF = Local Outlier Factor   Z   = Statistical Z-Score
ML_FEATURE_TABLE_DATA = [
    (
        "Payment amount (log scale)",
        "Natural log of the payment amount, compressing the wide range of amounts to a proportional scale.",
        "Continuous — no binary threshold",
        "IF, LOF",
        "Without log-scaling, a single extreme amount would dominate distance calculations in IF and LOF. "
        "The log transformation ensures amount magnitude is weighted proportionally rather than absolutely.",
    ),
    (
        "Amount vs. vendor average",
        "How much the payment amount differs from what this vendor is typically paid.",
        "Z-score > 2.0",
        "IF, LOF, Z-score",
        "Unusually large payments to a vendor may indicate over-billing or fictitious invoices.",
    ),
    (
        "Amount vs. cost centre average",
        "How much the payment amount differs from the typical amounts processed in that cost centre.",
        "Z-score > 2.0",
        "IF, LOF, Z-score",
        "Helps detect amounts that are out of place for the department, suggesting possible miscoding or inflated claims.",
    ),
    (
        "Round number",
        "Whether the payment amount ends in 00, 000, or 0,000.",
        "Exactly divisible by 100",
        "IF, LOF",
        "Genuine invoice amounts rarely end in round numbers; manually chosen or fictitious amounts often do.",
    ),
    (
        "Weekend payment",
        "Whether the voucher accounting date falls on a Saturday or Sunday.",
        "Saturday or Sunday",
        "IF, LOF",
        "Payments processed outside business hours may bypass the normal multi-person review and approval process.",
    ),
    (
        "Near approval threshold",
        "Whether the amount falls within 5% below a common approval limit",
        "Within 5% below SGD 1K / 5K / 10K / 50K / 100K",
        "IF, LOF",
        "A known technique ('structuring') to avoid triggering higher-level approval requirements.",
    ),
    (
        "Individual payee",
        "Whether the Vendor ID matches the Singapore NRIC/FIN format (one letter, 7 digits, one letter).",
        "Regex: ^[A-Za-z][0-9]{7}[A-Za-z]$",
        "IF, LOF",
        "Payments to individuals carry higher inherent risk versus registered businesses.",
    ),
    (
        "Processing time",
        "Number of calendar days between Invoice Date and Voucher Accounting Date.",
        "Absolute z-score > 2.5",
        "IF, LOF",
        "Very fast processing may indicate bypassed controls; unusually long delays may indicate backdating.",
    ),
    (
        "Description length",
        "Character length of the Voucher Line Description field.",
        "Absolute z-score > 2.5",
        "IF, LOF",
        "Very short descriptions may indicate incomplete entries; very long ones may indicate unusual or fabricated narrative.",
    ),
    (
        "Irregular repeated amount",
        "Same vendor paid the same amount more than twice, with no regular monthly/quarterly/annual schedule.",
        "> 2 occurrences with no detected recurring cycle",
        "IF, LOF",
        "May indicate duplicated or split payments that were structured to avoid detection.",
    ),
    (
        "Vendor billing consistency (CV)",
        "Coefficient of variation (std ÷ mean) of a vendor's positive payment amounts — how consistently "
        "they bill. High CV vendors have wide natural variance, making individual overpayments harder to "
        "detect via z-score alone.",
        "Continuous (higher = more variable)",
        "IF, LOF",
        "The vendor z-score compares each payment against the vendor's own billing history. "
        "For vendors with highly variable invoices, the standard deviation is naturally wide, making it harder "
        "for suspicious payments to breach the z-score threshold. The CV feature measures this variability — a high CV "
        "tells the ML models that the z-score signal is less reliable for that vendor, prompting greater weight to be placed on other features. ",
    ),
    (
        "Vendor transaction count",
        "Total number of payment lines attributed to this vendor in the dataset.",
        "Continuous — no binary threshold",
        "IF, LOF",
        "Vendors appearing very rarely (potentially fictitious) or at abnormally high volume relative "
        "to peers are a recognised fraud pattern; the ML models see this as contextual anomaly evidence.",
    ),
    (
        "Duplicate payment",
        "Whether the same invoice number, invoice date, vendor, and amount appears across more "
        "than one distinct payment voucher in the dataset.",
        "Same (Vendor ID, Invoice Number, Invoice Date, Amount) in > 1 Voucher ID",
        "IF, LOF",
        "Potential double payment of the same invoice — a common control failure or deliberate fraud technique.",
    ),
    (
        "Reversal / credit note",
        "Whether the payment amount is negative, indicating a reversal or credit note.",
        "Amount < 0",
        "IF, LOF",
        "Reversals paired with other risk signals on the corresponding original payment warrant auditor review.",
    ),
    (
        "Split purchase risk",
        "Whether the same vendor has two or more invoices on the same invoice date with alphanumerically "
        "sequential invoice number suffixes (e.g. INV-1001, INV-1002, INV-1003) AND the group total "
        "falls within 5% below $6,000 or $90,000 ($5,700–<$6,000 or $85,500–<$90,000).",
        "≥ 2 invoices from same vendor on same date with consecutive numeric suffixes; group total "
        "within $5,700–<$6,000 or $85,500–<$90,000",
        "IF, LOF",
        "A known technique for circumventing approval thresholds by splitting a single purchase into "
        "multiple invoices, each below the limit requiring higher-level authorisation.",
    ),
    (
        "Transposed amount",
        "Whether the same vendor and description group contains another transaction whose amount "
        "differs by exactly one pair of transposed digits in the whole-dollar portion "
        "(e.g. SGD 4,800 vs SGD 8,400). Cents are ignored, so cross-decimal swaps are not flagged.",
        "Exactly two digit positions swapped in the whole-dollar integer string (str(int(abs(amount)))) "
        "within same vendor + description group (positive amounts only)",
        "IF, LOF",
        "Digit transpositions are a common keying error that can result in significant over- or "
        "under-payment if left undetected. They are also occasionally used to conceal deliberate "
        "manipulation of invoice amounts.",
    ),
]

# Features that contribute to scoring but do not feed into any ML model.
BENFORD_FEATURE_TABLE_DATA = [
    (
        "Benford's Law first digit",
        "Whether the payment amount's first digit deviates significantly from Benford's expected frequency.",
        "First digit among the top-3 most deviant digits; non-recurring payments only.",
        "None — Benford's Law analysis only (5% of composite score)",
        "Systematic deviation may indicate manually constructed or manipulated amounts.",
    ),
]


def _render_feature_table(doc, data, col_widths):
    """Render a 5-column feature table (Feature | Measures | Threshold | ML Models | Why)."""
    headers = ["Feature", "What It Measures", "Threshold for Flagging", "ML Models", "Why It Matters"]
    tbl = doc.add_table(rows=1 + len(data), cols=5)
    tbl.style = 'Table Grid'

    hdr = tbl.rows[0]
    for i, (header, width) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _shade_cell(cell, "1F3864")
        cell.width = width

    for row_idx, row_data in enumerate(data, start=1):
        row   = tbl.rows[row_idx]
        shade = "F2F6FC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, (value, width) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[col_idx]
            cell.text = value
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _shade_cell(cell, shade)
            cell.width = width


def _page6_feature_table(doc):
    section = doc.add_section()
    _set_landscape(section)

    _heading(doc, "Feature Reference Table", level=1)
    _body(doc,
          "The tables below list each analytical feature, the threshold that determines whether a "
          "transaction is flagged, which ML scoring models the feature feeds into, and the audit "
          "rationale. ML Models: IF = Isolation Forest, LOF = Local Outlier Factor, "
          "Z-score = Statistical Z-Score Analysis.",
          size=9)
    doc.add_paragraph()

    col_widths = [Inches(1.6), Inches(2.1), Inches(1.8), Inches(1.2), Inches(3.8)]

    _heading(doc, "Features Used in Machine Learning Models", level=2)
    _body(doc,
          "The sixteen features below are candidates for the ML models in each run. Before fitting, "
          "Spearman correlation pruning removes one of any pair with |correlation| > 0.85, so the "
          "active feature set may be smaller than sixteen depending on the dataset. Surviving features "
          "are normalised via RobustScaler before being fed into the models. Amount z-scores "
          "additionally drive the Statistical Z-Score component directly.",
          size=9)
    doc.add_paragraph()
    _render_feature_table(doc, ML_FEATURE_TABLE_DATA, col_widths)

    doc.add_paragraph()

    _heading(doc, "Features Outside Machine Learning Models", level=2)
    _body(doc,
          "The feature below is computed by an independent method and contributes 5% of the "
          "composite risk score separately from the ML models.",
          size=9)
    doc.add_paragraph()
    _render_feature_table(doc, BENFORD_FEATURE_TABLE_DATA, col_widths)

    doc.add_paragraph()
    _body(doc,
          "References: Nigrini, M.J. (2012). Benford's Law: Applications for Forensic Accounting, "
          "Auditing, and Fraud Detection. ACFE Fraud Examiners Manual (current edition).",
          italic=True, size=7.5)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def export_word_report(df, df_vouchers, selected_vouchers, benford_stats, output_path,
                        excluded_count=0):
    print("  Building Word report...")
    doc = Document()

    section0 = doc.sections[0]
    section0.page_width    = Cm(21.0)
    section0.page_height   = Cm(29.7)
    section0.left_margin   = Cm(2.5)
    section0.right_margin  = Cm(2.5)
    section0.top_margin    = Cm(2.5)
    section0.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = 'Times New Roman'

    for st_name in ('Heading 1', 'Heading 2', 'List Bullet'):
        try:
            doc.styles[st_name].font.name = 'Times New Roman'
        except KeyError:
            pass

    print("    Page 1 — Scope and Limitations")
    _page_caveats(doc)
    doc.add_page_break()

    print("    Page 2 — Executive Summary")
    _page1(doc, df, df_vouchers, selected_vouchers, benford_stats)
    doc.add_page_break()

    print("    Page 3 — Methodology")
    _page2(doc, excluded_count)

    print("    Page 4 — Feature Reference Table")
    _page6_feature_table(doc)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"  Word report saved: {output_path}")
